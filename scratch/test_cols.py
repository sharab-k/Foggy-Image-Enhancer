import docx
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_two_columns(section):
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if not cols:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    else:
        cols = cols[0]
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '708')

doc = docx.Document()
doc.add_paragraph("This is the title", style='Title')
new_section = doc.add_section(WD_SECTION.CONTINUOUS)
set_two_columns(new_section)
for _ in range(50):
    doc.add_paragraph("This is a paragraph in the two column section. " * 10)
doc.save("test_cols.docx")
print("done")
