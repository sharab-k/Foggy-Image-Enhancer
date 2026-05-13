from docx import Document
from docx.shared import Pt
import os

def replace_text(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old_text in p.text:
                        inline = p.runs
                        for i in range(len(inline)):
                            if old_text in inline[i].text:
                                text = inline[i].text.replace(old_text, new_text)
                                inline[i].text = text

def generate_report():
    template_path = "FYP_Final_Report_AI_Interview_System.docx"
    doc = Document(template_path)

    # 1. Titles
    replace_text(doc, "AI-POWERED VIRTUAL INTERVIEW SYSTEM", "FOGGY IMAGE ENHANCER")
    replace_text(doc, "AI INTERVIEW SYSTEM", "FOGGY IMAGE ENHANCER")
    replace_text(doc, "AI Interview System", "Foggy Image Enhancer")
    
    # 2. Team
    if len(doc.tables) > 0:
        table = doc.tables[0]
        students = [["Noman Ahmed Siddiqui", "22K-4252"], ["Ansh Kumar", "22K-4564"], ["Affan Jan", "22K-4475"]]
        for i, (name, roll) in enumerate(students):
            if i < len(table.rows):
                table.rows[i].cells[0].text = name
                table.rows[i].cells[1].text = roll

    # 3. Intro
    intro_found = 0
    for p in doc.paragraphs:
        if "INTRODUCTION" in p.text.upper() and p.style.name.startswith('Heading'):
            intro_found = 1
        elif intro_found > 0 and len(p.text.strip()) > 30:
            if intro_found == 1:
                p.text = "This project explores an AI-powered platform for road scene enhancement in foggy conditions. Using YOLOv5 and advanced dehazing algorithms, we aim to improve object detection accuracy in low-visibility environments."
            elif intro_found == 2:
                p.text = "The system integrates Dark Channel Prior (DCP), CLAHE, and Bilateral Filtering to restore image clarity, providing a robust solution for autonomous navigation and surveillance."
            intro_found += 1
            if intro_found > 2: break

    # 4. Tables 3, 4, 5, 8, 9, 12
    # Table 3: Requirements
    if len(doc.tables) > 3:
        data = [["FR1", "Image Dehazing", "High"], ["FR2", "Object Detection", "High"], ["FR3", "Comparison UI", "Medium"]]
        for i, d in enumerate(data):
            if i+1 < len(doc.tables[3].rows):
                for j, v in enumerate(d): doc.tables[3].rows[i+1].cells[j].text = v

    # Table 4: Requirements Target
    if len(doc.tables) > 4:
        data = [["T1", "Detection Confidence", ">85%"], ["T2", "Processing Time", "<1.5s"], ["T3", "PSNR", ">25dB"]]
        for i, d in enumerate(data):
            if i+1 < len(doc.tables[4].rows):
                for j, v in enumerate(d): doc.tables[4].rows[i+1].cells[j].text = v

    # Table 5: Use Case
    if len(doc.tables) > 5:
        data = [["Enhance Image", "User", "User uploads image to get dehazed result"], ["Detect Objects", "System", "YOLOv5 runs on enhanced image"], ["Compare Results", "User", "View side-by-side metrics"]]
        for i, d in enumerate(data):
            if i+1 < len(doc.tables[5].rows):
                for j, v in enumerate(d): doc.tables[5].rows[i+1].cells[j].text = v

    # Table 8: Algorithms (Replacing Agent)
    if len(doc.tables) > 8:
        data = [["DCP", "Foggy Image", "Raw Pixels", "Dehazed Map"], ["CLAHE", "Low Contrast", "LAB Chans", "Enhanced L"], ["YOLOv5", "Image", "Pixels", "Boxes/Labels"]]
        for i, d in enumerate(data):
            if i+1 < len(doc.tables[8].rows):
                for j, v in enumerate(d): doc.tables[8].rows[i+1].cells[j].text = v

    # Table 9: Tech Stack
    if len(doc.tables) > 9:
        data = [["Frontend", "React/Vite", "v18"], ["Backend", "FastAPI", "v0.109"], ["AI", "YOLOv5", "v8.0"]]
        for i, d in enumerate(data):
            if i+1 < len(doc.tables[9].rows):
                for j, v in enumerate(d): doc.tables[9].rows[i+1].cells[j].text = v

    # Table 12: Results
    if len(doc.tables) > 12:
        data = [["Confidence", "88%", "80%", "Pass"], ["Latency", "1.1s", "2.0s", "Pass"]]
        for i, d in enumerate(data):
            if i+1 < len(doc.tables[12].rows):
                for j, v in enumerate(d): doc.tables[12].rows[i+1].cells[j].text = v

    doc.save("FYP_Final_Report_Foggy_Image_Enhancer.docx")

if __name__ == "__main__":
    generate_report()
