from docx import Document

def read_template(file_path):
    doc = Document(file_path)
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading') or para.text.isupper():
            print(f"Para {i} ({para.style.name}): {para.text}")
        elif i < 50: # First few paras for title page
             print(f"Para {i} ({para.style.name}): {para.text}")
    
    for i, table in enumerate(doc.tables):
        print(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} columns")
        # Print first row of each table to understand context
        if len(table.rows) > 0:
            print(f"  Header: {[cell.text.strip() for cell in table.rows[0].cells]}")

if __name__ == "__main__":
    read_template("FYP_Final_Report_AI_Interview_System.docx")
