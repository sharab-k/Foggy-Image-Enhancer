import docx

def analyze_docx(filepath):
    try:
        doc = docx.Document(filepath)
        print("=== PARAGRAPHS ===")
        for i, p in enumerate(doc.paragraphs):
            if p.style.name.startswith('Heading'):
                print(f"[{p.style.name}] {p.text}")
            elif i < 30: # Print first 30 paragraphs to see title page structure
                if p.text.strip():
                    print(f"[Style: {p.style.name}] {p.text}")
        
        print("\n=== TABLES ===")
        for i, table in enumerate(doc.tables):
            print(f"Table {i+1}: {len(table.rows)} rows, {len(table.columns)} cols, Style: {table.style.name}")
            # Print header row if possible
            try:
                header = [cell.text.strip() for cell in table.rows[0].cells]
                print(f"  Header: {header}")
            except:
                pass
    except Exception as e:
        print(f"Error: {e}")

if __name__ == '__main__':
    analyze_docx('FYP_Final_Report_AI_Interview_System.docx')
