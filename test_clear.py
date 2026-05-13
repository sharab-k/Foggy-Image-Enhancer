import docx

def clear_doc(filepath):
    doc = docx.Document(filepath)
    # delete all body elements to retain styles/header/footer
    for element in doc.element.body:
        doc.element.body.remove(element)
    
    doc.add_paragraph("Test Title", style='Normal')
    doc.add_heading("1. Introduction", level=1)
    doc.save("test_clear.docx")
    print("Done")

if __name__ == '__main__':
    clear_doc('FYP_Final_Report_AI_Interview_System.docx')
