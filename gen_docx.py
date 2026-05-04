from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def main():
    doc = Document()

    # Style modifications to match standard reports
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    style = doc.styles['Heading 1']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.bold = True
    
    style = doc.styles['Heading 2']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.bold = True

    # CS4092 Final Year Project
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("CS4092\n").bold = False
    p.add_run("Final Year Project\n").bold = False
    run = p.add_run("Foggy Image Enhancer")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph("Submitted By")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    # Table for Students
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Student Name'
    hdr_cells[1].text = 'Roll Number'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Noman Ahmed Siddiqui'
    row_cells[1].text = '22K-4252'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'Ansh Kumar'
    row_cells[1].text = '22K-4564'
    
    row_cells = table.rows[3].cells
    row_cells[0].text = 'Affan Jan'
    row_cells[1].text = '22K-4475'

    # Add space
    doc.add_paragraph()

    # Project Progress Report Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Project Progress Report\n").bold = True
    p.add_run("Department of Computer Science\n")
    p.add_run("National University of Computer and Emerging Sciences\n")
    p.add_run("FAST – Karachi Campus").bold = False
    
    doc.add_heading("1  INTRODUCTION", level=1)
    doc.add_paragraph(
        "This project explores an AI-powered computer vision platform designed to overcome the challenge of object detection in poor visibility. Specifically, it focuses on enhancing road scene visibility under foggy conditions and analyzes how these enhancements impact YOLOv5 object detection performance. The system allows users to upload blurry or foggy images, automatically applies enhancement algorithms, and then runs real-time object detection natively on both the original and enhanced images for comparative analysis."
    )
    doc.add_paragraph(
        "The image enhancement pipeline seamlessly integrates techniques such as Contrast Limited Adaptive Histogram Equalization (CLAHE), Gamma Correction, Bilateral Filtering, and Dark Channel Prior (DCP) to restore visibility. The YOLOv5 object detection model then evaluates the processed outputs against the originals to demonstrate empirical improvements in accuracy and bounding box confidence scores. During FYP Part 1, the core image enhancement algorithms were developed, object detection models were benchmarked, and the initial backend API was built. FYP Part 2 focuses heavily on developing an interactive React frontend, containerizing the application using Docker, optimizing the heavy computer vision operations for resource-constrained free-tier cloud environments (Hugging Face Spaces and Vercel), and resolving cross-origin deployment errors."
    )

    doc.add_heading("2  TIMELINE", level=1)
    doc.add_paragraph("The following timeline was established during FYP Part 1 and covers the work completed in that phase.")
    p = doc.add_paragraph("FYP Part 1 Timeline (Completed)")
    p.runs[0].bold = True

    tab1 = doc.add_table(rows=6, cols=2)
    tab1.style = 'Table Grid'
    rc = tab1.rows[0].cells
    rc[0].text = 'Period'
    rc[1].text = 'Activities'
    for cell in rc:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    rc = tab1.rows[1].cells
    rc[0].text = 'Weeks 1–3'
    rc[1].text = 'Requirements gathering, literature review on dehazing algorithms, and system architecture design'
    rc = tab1.rows[2].cells
    rc[0].text = 'Weeks 4–6'
    rc[1].text = 'Implementation of core enhancement techniques (CLAHE, Gamma Correction, Bilateral Filtering, DCP)'
    rc = tab1.rows[3].cells
    rc[0].text = 'Weeks 7–9'
    rc[1].text = 'YOLOv5 object detection setup and baseline comparison evaluation on foggy datasets'
    rc = tab1.rows[4].cells
    rc[0].text = 'Weeks 10–12'
    rc[1].text = 'Backend development using FastAPI and encapsulation of the processing logic'
    rc = tab1.rows[5].cells
    rc[0].text = 'Weeks 13–15'
    rc[1].text = 'Pipeline testing and analysis of enhancement effects on object detection confidence scores'

    doc.add_heading("3  PROGRESS", level=1)
    doc.add_paragraph("FYP Part 2 began after the completion of the core computer vision algorithms and baseline application structure in FYP Part 1. The following subsections describe the work completed so far in FYP Part 2, leading up to the mid-evaluation.")

    doc.add_heading("3.1  FRONTEND DEVELOPMENT AND INTEGRATION", level=2)
    doc.add_paragraph("To provide an interactive platform for the models, a dynamic frontend interface was developed during FYP Part 2 using Vite and React. The application features a robust drag-and-drop file upload mechanism, enabling users to seamlessly upload their test images. The interface provides a side-by-side comparison view allowing direct visual analysis of the original image against the algorithmically enhanced image. Furthermore, it overlays YOLOv5 bounding boxes and confidence score labels to visually demonstrate the object detection pipeline’s outputs. All data transfers were optimized via REST API routes using efficient multipart form-data handling for processing image streams between the React frontend and FastAPI backend.")

    doc.add_heading("3.2  DEPLOYMENT AND CLOUD OPTIMIZATION", level=2)
    doc.add_paragraph("A substantial part of FYP Part 2 was dedicated to transitioning the project from a local development codebase to a fully deployed cloud infrastructure. Docker containerization was set up for the FastAPI backend to seamlessly host the inference models on Hugging Face Spaces. The React frontend was successfully deployed to Vercel. A significant challenge faced and resolved was optimizing the resource-heavy computer vision pipeline to execute reliably within the strict memory and compute limitations of Hugging Face’s free tier. Additionally, widespread Cross-Origin Resource Sharing (CORS) setup issues and module path errors were mitigated to guarantee stable and uninterrupted bidirectional communication between the Vercel frontend and Hugging Face backend.")

    p = doc.add_paragraph("FYP Part 2 – Work Completed (Before Mid Evaluation)")
    p.runs[0].bold = True

    tab2 = doc.add_table(rows=6, cols=3)
    tab2.style = 'Table Grid'
    rc = tab2.rows[0].cells
    rc[0].text = 'Task'
    rc[1].text = 'Description'
    rc[2].text = 'Status'
    for cell in rc:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    rc = tab2.rows[1].cells
    rc[0].text = 'Frontend Development'
    rc[1].text = 'Built Vite/React interface with side-by-side visual comparison view'
    rc[2].text = 'Completed'

    rc = tab2.rows[2].cells
    rc[0].text = 'Backend Integration'
    rc[1].text = 'Unified frontend and FastAPI backend APIs for image streaming'
    rc[2].text = 'Completed'

    rc = tab2.rows[3].cells
    rc[0].text = 'Cloud Containerization'
    rc[1].text = 'Created Docker configuration for seamless backend deployment'
    rc[2].text = 'Completed'

    rc = tab2.rows[4].cells
    rc[0].text = 'Pipeline Optimization'
    rc[1].text = 'Optimized YOLOv5 inferences to operate within Hugging Face free-tier constraints'
    rc[2].text = 'Completed'
    
    rc = tab2.rows[5].cells
    rc[0].text = 'CORS and Deployment'
    rc[1].text = 'Resolved cross-origin and Vercel build errors; achieved full deployment'
    rc[2].text = 'Completed'

    doc.add_paragraph()
    p = doc.add_paragraph("Remaining Milestones (After Mid Evaluation – 1 Month)")
    p.runs[0].bold = True
    doc.add_paragraph("The following three milestones are planned for the remaining one month of FYP Part 2, from mid evaluation to the final submission.")
    
    tab3 = doc.add_table(rows=4, cols=4)
    tab3.style = 'Table Grid'
    rc = tab3.rows[0].cells
    rc[0].text = '#'
    rc[1].text = 'Milestone'
    rc[2].text = 'Timeline'
    rc[3].text = 'Status'
    for cell in rc:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    rc = tab3.rows[1].cells
    rc[0].text = '1'
    rc[1].text = 'Advanced Analytics Metrics Comparison'
    rc[2].text = 'Weeks 1–2'
    rc[3].text = 'Planned'

    rc = tab3.rows[2].cells
    rc[0].text = '2'
    rc[1].text = 'End-to-End Latency Improvements'
    rc[2].text = 'Week 3'
    rc[3].text = 'Planned'

    rc = tab3.rows[3].cells
    rc[0].text = '3'
    rc[1].text = 'System Testing and Documentation'
    rc[2].text = 'Weeks 3–4'
    rc[3].text = 'Planned'

    doc.add_paragraph()
    doc.add_paragraph("Milestone 1 involves generating analytical metrics within the React frontend, automatically comparing detection confidence percentages before and after enhancements.")
    doc.add_paragraph("Milestone 2 addresses end-to-end latency optimizations such as payload compression and response buffering to ensure the YOLOv5 and DCP algorithms remain responsive.")
    doc.add_paragraph("Milestone 3 covers final validation, load testing of the cloud endpoints, bug fixing, and developing the final project documentation and manuals.")

    doc.add_heading("4  UPDATED TIMELINE", level=1)
    doc.add_paragraph("The following updated timeline covers the entire project from proposal to completion, spanning both FYP Part 1 and FYP Part 2.")

    tab4 = doc.add_table(rows=12, cols=3)
    tab4.style = 'Table Grid'
    rc = tab4.rows[0].cells
    rc[0].text = 'Phase'
    rc[1].text = 'Period'
    rc[2].text = 'Activities'
    for cell in rc:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
    
    rc = tab4.rows[1].cells
    rc[0].text = 'FYP 1'
    rc[1].text = 'Weeks 1–3'
    rc[2].text = 'Requirements, architecture, literature review, and algorithm design'
    
    rc = tab4.rows[2].cells
    rc[0].text = 'FYP 1'
    rc[1].text = 'Weeks 4–6'
    rc[2].text = 'Implementation of enhancement algorithms (CLAHE, DCP, Gamma, etc.)'
    
    rc = tab4.rows[3].cells
    rc[0].text = 'FYP 1'
    rc[1].text = 'Weeks 7–9'
    rc[2].text = 'YOLOv5 integration and baseline image fog evaluations'
    
    rc = tab4.rows[4].cells
    rc[0].text = 'FYP 1'
    rc[1].text = 'Weeks 10–12'
    rc[2].text = 'Backend FastAPI development and container logic mock-up'
    
    rc = tab4.rows[5].cells
    rc[0].text = 'FYP 1'
    rc[1].text = 'Weeks 13–15'
    rc[2].text = 'Pipeline design and initial qualitative testing'
    
    rc = tab4.rows[6].cells
    rc[0].text = 'FYP 2'
    rc[1].text = 'Weeks 1–3'
    rc[2].text = 'Frontend development (Vite + React framework) (Completed)'
    
    rc = tab4.rows[7].cells
    rc[0].text = 'FYP 2'
    rc[1].text = 'Weeks 4–6'
    rc[2].text = 'Frontend & Backend integration, REST endpoint stability (Completed)'
    
    rc = tab4.rows[8].cells
    rc[0].text = 'FYP 2'
    rc[1].text = 'Weeks 7–8'
    rc[2].text = 'Dockerize backend, Hugging Face deployment, optimization (Completed)'
    
    rc = tab4.rows[9].cells
    rc[0].text = 'FYP 2'
    rc[1].text = 'Weeks 9–10'
    rc[2].text = 'Advanced analytics metric integration UI (Planned)'
    
    rc = tab4.rows[10].cells
    rc[0].text = 'FYP 2'
    rc[1].text = 'Week 11'
    rc[2].text = 'End-to-End inference latency improvements (Planned)'
    
    rc = tab4.rows[11].cells
    rc[0].text = 'FYP 2'
    rc[1].text = 'Weeks 11–12'
    rc[2].text = 'Stress testing, code refactor, and final reporting (Planned)'

    doc.add_heading("REFERENCES", level=1)
    doc.add_paragraph("[1] K. He, J. Sun, and X. Tang, “Single Image Haze Removal Using Dark Channel Prior,” IEEE Transactions on Pattern Analysis and Machine Intelligence, 2011.")
    doc.add_paragraph("[2] J. Redmon et al., “You Only Look Once: Unified, Real-Time Object Detection,” arXiv, 2015.")
    doc.add_paragraph("[3] Ultralytics, “YOLOv5 Documentation,” 2020.")
    doc.add_paragraph("[4] FastAPI Documentation, https://fastapi.tiangolo.com.")
    doc.add_paragraph("[5] Hugging Face Documentation, https://huggingface.co/docs.")
    doc.add_paragraph("[6] Vite Documentation, https://vitejs.dev/guide/.")
    doc.add_paragraph("[7] Docker Documentation, https://docs.docker.com.")

    doc.save("FYP_II_Progress_Report_Foggy_Image_Enhancer.docx")

if __name__ == "__main__":
    main()
